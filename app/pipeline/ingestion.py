from __future__ import annotations

import base64
import io
from typing import List, Dict, Any
from pathlib import Path

import pandas as pd
from PIL import Image

from app.models.schemas import UploadDocument
from app.models.config import get_settings
from app.retrieval.service import HybridRetrievalService
from app.intelligence.table_extraction import TableExtractor, FigureExtractor, DocumentStructureAnalyzer
from app.intelligence.multilang import LanguageDetector, DocumentLanguageAnalyzer
from app.intelligence.versioning import DocumentVersionManager
from app.intelligence.semantic_chunking import SemanticChunker, ChunkingConfig
from app.intelligence.math_extraction import MathFormulaExtractor
from app.intelligence.code_extraction import CodeSnippetExtractor


class IngestionPipeline:
    def __init__(self, retriever: HybridRetrievalService | None = None) -> None:
        self.retriever = retriever or HybridRetrievalService()
        self.settings = get_settings()
        
        # Initialize advanced intelligence components
        self.table_extractor = TableExtractor()
        self.figure_extractor = FigureExtractor()
        self.structure_analyzer = DocumentStructureAnalyzer()
        self.language_detector = LanguageDetector()
        self.language_analyzer = DocumentLanguageAnalyzer()
        self.version_manager = DocumentVersionManager()
        self.semantic_chunker = SemanticChunker(ChunkingConfig(
            min_chunk_size=200,
            max_chunk_size=1000,
            similarity_threshold=0.3
        ))
        self.math_extractor = MathFormulaExtractor()
        self.code_extractor = CodeSnippetExtractor()
        
        # Ensure uploads dir exists
        self.uploads_dir = Path(self.settings.uploads_dir)
        self.uploads_dir.mkdir(parents=True, exist_ok=True)

    def _decode_content(self, doc: UploadDocument) -> bytes:
        if doc.content_base64:
            return base64.b64decode(doc.content_base64)
        raise ValueError("URL fetch not implemented in this skeleton")

    def _extract_text_from_pdf(self, content: bytes) -> str:
        try:
            import pypdf

            reader = pypdf.PdfReader(io.BytesIO(content))
            texts = []
            for page in reader.pages:
                texts.append(page.extract_text() or "")
            candidate = "\n".join(texts)
            if candidate.strip():
                return candidate
        except Exception:
            pass
        # Fallback to OCR for scanned PDFs
        try:
            from pdf2image import convert_from_bytes
            import pytesseract

            images = convert_from_bytes(content)
            ocr_texts = []
            for img in images:
                ocr_texts.append(pytesseract.image_to_string(img))
            return "\n".join(ocr_texts)
        except Exception:
            return ""

    def _extract_texts_from_pdf(self, content: bytes) -> List[str]:
        # Page-wise extraction; falls back to OCR page-wise
        try:
            import pypdf

            reader = pypdf.PdfReader(io.BytesIO(content))
            page_texts: List[str] = []
            for page in reader.pages:
                page_texts.append(page.extract_text() or "")
            if any(t.strip() for t in page_texts):
                return page_texts
        except Exception:
            pass
        try:
            from pdf2image import convert_from_bytes
            import pytesseract

            images = convert_from_bytes(content)
            page_texts = [pytesseract.image_to_string(img) for img in images]
            return page_texts
        except Exception:
            return []

    def _extract_text_from_docx(self, content: bytes) -> str:
        try:
            import docx

            document = docx.Document(io.BytesIO(content))
            return "\n".join([p.text for p in document.paragraphs])
        except Exception:
            return ""

    def _extract_text_from_xlsx(self, content: bytes) -> str:
        try:
            xls = pd.read_excel(io.BytesIO(content), sheet_name=None)
            texts = []
            for name, df in xls.items():
                texts.append(f"# Sheet: {name}\n{df.to_csv(index=False)}")
            return "\n\n".join(texts)
        except Exception:
            return ""

    def _extract_text_from_csv(self, content: bytes) -> str:
        try:
            df = pd.read_csv(io.BytesIO(content))
            return df.to_csv(index=False)
        except Exception:
            return content.decode("utf-8", errors="ignore")

    def _extract_text_from_image(self, content: bytes) -> str:
        try:
            import pytesseract

            img = Image.open(io.BytesIO(content))
            return pytesseract.image_to_string(img)
        except Exception:
            return ""

    def _extract_text_from_txt(self, content: bytes) -> str:
        return content.decode("utf-8", errors="ignore")

    def extract(self, doc: UploadDocument) -> Dict[str, Any]:
        """Enhanced extraction with advanced intelligence."""
        content = self._decode_content(doc)
        
        # Persist original file for audit/cache
        try:
            target = self.uploads_dir / doc.filename
            with target.open("wb") as f:
                f.write(content)
        except Exception:
            pass
        
        # Extract basic text content
        text_content = self._extract_basic_text(doc, content)
        
        # Initialize extraction result
        extraction_result = {
            'text': text_content,
            'tables': [],
            'figures': [],
            'structure': {},
            'language_info': {},
            'metadata': {
                'filename': doc.filename,
                'mime_type': doc.mime_type,
                'size_bytes': len(content)
            }
        }
        
        # Language detection and analysis
        if text_content:
            lang_info = self.language_detector.detect_language(text_content)
            lang_analysis = self.language_analyzer.analyze_document_languages(text_content)
            extraction_result['language_info'] = {
                'primary_language': lang_info.language,
                'confidence': lang_info.confidence,
                'script': lang_info.script,
                'analysis': lang_analysis
            }
        
        # Document structure analysis
        if text_content:
            structure = self.structure_analyzer.analyze_structure(text_content, doc.filename)
            extraction_result['structure'] = structure
        
        # Advanced extraction for PDFs
        mt = doc.mime_type.lower()
        if mt in {"application/pdf"}:
            # Extract tables
            try:
                tables = self.table_extractor.extract_tables_from_pdf(content, doc.filename)
                extraction_result['tables'] = tables
            except Exception as e:
                print(f"Table extraction failed: {e}")
            
            # Extract figures
            try:
                figures = self.figure_extractor.extract_figures_from_pdf(content, doc.filename)
                extraction_result['figures'] = figures
            except Exception as e:
                print(f"Figure extraction failed: {e}")
        
        # Mathematical formula extraction
        if text_content:
            try:
                math_result = self.math_extractor.extract_formulas_from_text(text_content, doc.filename)
                extraction_result['math_formulas'] = {
                    'total_formulas': math_result.total_formulas,
                    'formulas_by_type': math_result.formulas_by_type,
                    'formulas': [
                        {
                            'id': f.formula_id,
                            'latex': f.latex_code,
                            'type': f.formula_type,
                            'complexity': f.complexity_score,
                            'variables': f.variables,
                            'confidence': f.confidence
                        } for f in math_result.extracted_formulas
                    ]
                }
            except Exception as e:
                print(f"Math extraction failed: {e}")
                extraction_result['math_formulas'] = {'total_formulas': 0, 'formulas': []}
        
        # Code snippet extraction
        if text_content:
            try:
                code_result = self.code_extractor.extract_code_from_text(text_content, doc.filename)
                extraction_result['code_snippets'] = {
                    'total_snippets': code_result.total_snippets,
                    'snippets_by_language': code_result.snippets_by_language,
                    'snippets': [
                        {
                            'id': s.snippet_id,
                            'language': s.language,
                            'confidence': s.confidence,
                            'complexity': s.complexity_score,
                            'functions': s.functions,
                            'classes': s.classes,
                            'keywords': s.keywords[:10],  # Limit for storage
                            'syntax_valid': s.syntax_valid
                        } for s in code_result.extracted_snippets
                    ]
                }
            except Exception as e:
                print(f"Code extraction failed: {e}")
                extraction_result['code_snippets'] = {'total_snippets': 0, 'snippets': []}
        
        # Version tracking
        try:
            version = self.version_manager.add_version(
                document_id=doc.filename,
                filename=doc.filename,
                content=text_content,
                metadata=extraction_result['metadata']
            )
            extraction_result['version_info'] = {
                'version_id': version.version_id,
                'is_new_version': True,
                'change_summary': version.change_summary
            }
        except Exception as e:
            print(f"Version tracking failed: {e}")
        
        return extraction_result
    
    def _extract_basic_text(self, doc: UploadDocument, content: bytes) -> str:
        """Extract basic text content from document."""
        mt = doc.mime_type.lower()
        if mt in {"application/pdf"}:
            return self._extract_text_from_pdf(content)
        if mt in {"application/vnd.openxmlformats-officedocument.wordprocessingml.document"}:
            return self._extract_text_from_docx(content)
        if mt in {"application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"}:
            return self._extract_text_from_xlsx(content)
        if mt in {"text/csv"}:
            return self._extract_text_from_csv(content)
        if mt.startswith("image/"):
            return self._extract_text_from_image(content)
        if mt in {"text/plain"}:
            return self._extract_text_from_txt(content)
        return ""

    def chunk(self, text: str, filename: str = "") -> List[str]:
        """Enhanced chunking using semantic boundaries."""
        try:
            # Use semantic chunking
            semantic_chunks = self.semantic_chunker.chunk_document(text, filename)
            return [chunk.text for chunk in semantic_chunks]
        except Exception as e:
            print(f"Semantic chunking failed, falling back to simple chunking: {e}")
            # Fallback to simple chunking
            chunks = [c.strip() for c in text.split("\n\n") if c.strip()]
            return chunks

    def process(self, documents: List[UploadDocument], overwrite: bool = False) -> int:
        """Enhanced processing with advanced intelligence features."""
        ids: List[str] = []
        texts: List[str] = []
        metadatas: List[dict] = []
        
        for doc in documents:
            base_id = doc.id or doc.filename
            
            # Use enhanced extraction
            extraction_result = self.extract(doc)
            text_content = extraction_result['text']
            
            if not text_content:
                continue
            
            # Enhanced metadata from extraction
            base_metadata = {
                "source": doc.filename,
                "mime_type": doc.mime_type,
                "language": extraction_result.get('language_info', {}).get('primary_language', 'en'),
                "language_confidence": extraction_result.get('language_info', {}).get('confidence', 0.5),
                "has_tables": len(extraction_result.get('tables', [])) > 0,
                "has_figures": len(extraction_result.get('figures', [])) > 0,
                "num_sections": len(extraction_result.get('structure', {}).get('sections', [])),
                "version_id": extraction_result.get('version_info', {}).get('version_id'),
            }
            
            # Chunk the text content
            chunks = self.chunk(text_content, doc.filename)
            
            for ci, chunk_text in enumerate(chunks):
                chunk_id = f"{base_id}__c{ci:04d}"
                
                # Create enhanced metadata for this chunk
                chunk_metadata = base_metadata.copy()
                chunk_metadata.update({
                    "chunk_index": ci,
                    "chunk_type": "content"
                })
                
                ids.append(chunk_id)
                texts.append(chunk_text)
                metadatas.append(chunk_metadata)
            
            # Index table content separately if available
            for table_idx, table in enumerate(extraction_result.get('tables', [])):
                table_id = f"{base_id}__table_{table_idx:04d}"
                table_metadata = base_metadata.copy()
                table_metadata.update({
                    "chunk_type": "table",
                    "table_method": table.get('method'),
                    "table_confidence": table.get('confidence', 0.5),
                    "table_shape": str(table.get('shape')),
                    "page": table.get('page')
                })
                
                # Use CSV representation of table as text
                table_text = table.get('csv', '')
                if table_text:
                    ids.append(table_id)
                    texts.append(f"Table data:\n{table_text}")
                    metadatas.append(table_metadata)
            
            # Index figure content if available
            for fig_idx, figure in enumerate(extraction_result.get('figures', [])):
                figure_id = f"{base_id}__figure_{fig_idx:04d}"
                figure_metadata = base_metadata.copy()
                figure_metadata.update({
                    "chunk_type": "figure",
                    "figure_type": figure.get('type'),
                    "figure_confidence": figure.get('confidence', 0.5),
                    "page": figure.get('page')
                })
                
                # Use figure description as text
                figure_text = f"Figure: {figure.get('caption', '')}\nDescription: {figure.get('description', '')}"
                if figure_text.strip():
                    ids.append(figure_id)
                    texts.append(figure_text)
                    metadatas.append(figure_metadata)
        
        # Index all content
        if ids:
            self.retriever.index(ids, texts, metadatas)
        
        return len(ids)

    def ingest_from_directory(self, directory: str) -> int:
        from app.models.schemas import UploadDocument
        import mimetypes
        import base64

        dir_path = Path(directory)
        files = [p for p in dir_path.glob("**/*") if p.is_file()]
        docs: List[UploadDocument] = []
        for p in files:
            mime, _ = mimetypes.guess_type(str(p))
            if not mime:
                continue
            if not (mime.startswith("image/") or mime in {
                "application/pdf",
                "text/plain",
                "text/csv",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            }):
                continue
            content = p.read_bytes()
            docs.append(
                UploadDocument(
                    filename=p.name,
                    mime_type=mime,
                    content_base64=base64.b64encode(content).decode("utf-8"),
                )
            )
        return self.process(docs, overwrite=False)


