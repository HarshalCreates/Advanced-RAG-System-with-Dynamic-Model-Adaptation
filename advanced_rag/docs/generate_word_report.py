from __future__ import annotations

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime


TITLE = "Advanced Multi-Modal RAG System – Project Documentation"
OUTPUT_PATH = "docs/Advanced_RAG_Project_Documentation.docx"


def add_heading(document: Document, text: str, level: int = 1) -> None:
    document.add_heading(text, level=level)


def add_paragraph(document: Document, text: str) -> None:
    document.add_paragraph(text)


def add_bullet(document: Document, text: str) -> None:
    p = document.add_paragraph(style="List Bullet")
    p.add_run(text)


def add_numbered(document: Document, text: str) -> None:
    p = document.add_paragraph(style="List Number")
    p.add_run(text)


def add_code_block(document: Document, text: str) -> None:
    # Basic monospaced block using a normal paragraph with font override
    p = document.add_paragraph()
    run = p.add_run(text)
    font = run.font
    font.name = "Courier New"
    font.size = Pt(10)


def build_doc() -> Document:
    doc = Document()

    # Title
    title = doc.add_paragraph()
    run = title.add_run(TITLE)
    run.bold = True
    run.font.size = Pt(18)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_paragraph(doc, f"Generated on {datetime.utcnow().strftime('%Y-%m-%d %H:%M:%SZ')} (UTC)")

    # 1. Overview
    add_heading(doc, "1. Overview", 1)
    add_paragraph(
        doc,
        "This project implements a production-grade, multi-modal Retrieval-Augmented Generation (RAG) system "
        "with dynamic model adaptation. It supports ingestion of heterogeneous documents (PDF, DOCX, XLSX, CSV, TXT, "
        "and images via OCR), hybrid retrieval (dense, TF-IDF, BM25; fusion ranking), and a FastAPI backend with REST and WebSockets. "
        "An admin API enables live hot-swapping of embedding/generation/retriever backends. A Chainlit UI provides chat with real-time model switching.",
    )

    # 2. Stack
    add_heading(doc, "2. Frameworks & Stack", 1)
    add_bullet(doc, "Language: Python 3.11+")
    add_bullet(doc, "API: FastAPI + Uvicorn")
    add_bullet(doc, "UI: Chainlit (chat UI with settings)")
    add_bullet(doc, "Vector: FAISS (CPU) / In-memory; BM25 (rank-bm25); TF-IDF (scikit-learn)")
    add_bullet(doc, "Embeddings: OpenAI / Cohere / SentenceTransformers / Hash fallback")
    add_bullet(doc, "Generation: GPT-4o (OpenAI) / Claude (Anthropic) / Llama (Ollama) / Echo fallback")
    add_bullet(doc, "OCR: Tesseract; Scanned PDF OCR via pdf2image + Tesseract")
    add_bullet(doc, "Observability: Prometheus metrics; Structured logging")

    # 3. Project Structure
    add_heading(doc, "3. Project Structure", 1)
    add_code_block(
        doc,
        """
advanced_rag/
├── app/
│   ├── main.py                      # FastAPI app factory & wiring
│   ├── api/                         # REST & WebSocket routes
│   ├── models/                      # Pydantic schemas & settings
│   ├── pipeline/                    # Ingestion & query pipeline manager
│   ├── ingestion/                   # Structure-aware chunker (extensible)
│   ├── retrieval/                   # Vector backends, sparse, BM25, fusion, service
│   ├── index/                       # Document store (persistent)
│   ├── embeddings/                  # Embedding clients & factory
│   ├── generation/                  # Generation clients & factory
│   ├── admin/                       # Admin routes & demo, dashboard HTML
│   ├── observability/               # Logging & metrics
│   └── security/                    # Security middleware
├── ui/chainlit/                     # Chainlit UI app
├── docs/                            # Documentation & this generator
├── tests/                           # Basic tests
├── requirements.txt                 # Python dependencies
├── Dockerfile, docker-compose.yml   # Containerization
└── README.md                        # Quickstart (optional)
        """.strip(),
    )

    # 4. Setup & Running
    add_heading(doc, "4. Setup & Running", 1)
    add_paragraph(doc, "Local (venv):")
    add_code_block(
        doc,
        """
cd advanced_rag
python3 -m venv .venv
source .venv/bin/activate
pip install -r requirements.txt
uvicorn app.main:app --host 127.0.0.1 --port 8000 --reload
        """.strip(),
    )
    add_paragraph(doc, "Docker:")
    add_code_block(
        doc,
        """
docker compose up --build
        """.strip(),
    )
    add_paragraph(doc, "Chainlit UI:")
    add_code_block(
        doc,
        """
export RAG_API_BASE="http://127.0.0.1:8000"
chainlit run ui/chainlit/langchain_app.py -w --host 127.0.0.1 --port 8501
        """.strip(),
    )

    # 5. Configuration
    add_heading(doc, "5. Configuration", 1)
    add_paragraph(
        doc,
        "Settings are defined in app/models/config.py (pydantic-settings). You can override via environment variables or .env.",
    )
    add_bullet(doc, "environment: development | staging | production")
    add_bullet(doc, "log_level: DEBUG | INFO | WARN | ERROR")
    add_bullet(doc, "data_dir: ./data; index_dir: ./data/index; uploads_dir: ./data/uploads")
    add_bullet(doc, "embedding_backend: openai | cohere | sentence-transformers | hash")
    add_bullet(doc, "embedding_model: model string or 'hash'")
    add_bullet(doc, "generation_backend: openai | anthropic | ollama")
    add_bullet(doc, "generation_model: e.g., gpt-4o | claude-3-5-sonnet-20240620 | llama3")
    add_bullet(doc, "retriever_backend: faiss | chroma (in-memory stub)")
    add_bullet(doc, "openai_api_key, cohere_api_key, anthropic_api_key (optional)")
    add_bullet(doc, "admin_api_key: protect admin hot-swap routes (send in X-API-Key)")

    # 6. Endpoints
    add_heading(doc, "6. API Endpoints", 1)
    add_bullet(doc, "GET /api/health – Service health")
    add_bullet(doc, "POST /api/ingest – Ingest documents (JSON-encoded base64)")
    add_bullet(doc, "POST /api/ingest/upload – Upload files via multipart form")
    add_bullet(doc, "POST /api/ingest/directory – Ingest server-side directory")
    add_bullet(doc, "POST /api/query – Query; returns schema-compliant JSON")
    add_bullet(doc, "WS /ws/stream – Streaming tokens with confidence updates")
    add_bullet(doc, "GET /api/admin/dashboard – Admin UI (ingestion & hot-swap)")
    add_bullet(doc, "POST /api/admin/hot-swap/embeddings – Swap embedding backend+model")
    add_bullet(doc, "POST /api/admin/hot-swap/generation – Swap generation backend+model")
    add_bullet(doc, "POST /api/admin/hot-swap/retriever – Swap retriever backend")

    # 7. Ingestion Pipeline
    add_heading(doc, "7. Document Ingestion & Multi-Modal Extraction", 1)
    add_paragraph(
        doc,
        "Ingestion decodes content, persists files to ./data/uploads, extracts text depending on type, and performs "
        "structure-aware chunking. PDFs are processed page-wise preserving page numbers; scanned PDFs are OCRed via pdf2image + Tesseract.",
    )
    add_bullet(doc, "PDF: pypdf page text; OCR fallback per page")
    add_bullet(doc, "DOCX: python-docx paragraphs")
    add_bullet(doc, "XLSX: pandas read_excel; each sheet to CSV-like text")
    add_bullet(doc, "CSV/TXT: decoded text")
    add_bullet(doc, "Images: Tesseract OCR")
    add_paragraph(doc, "Chunk IDs encode provenance: filename__p0001__c0003 (PDF page/chunk).")

    # 8. Retrieval Architecture
    add_heading(doc, "8. Retrieval Architecture", 1)
    add_paragraph(
        doc,
        "Hybrid retrieval combines dense vector search with sparse methods (TF-IDF and BM25). Fusion ranking merges scores. "
        "Metadata filtering is supported. A graph reranker is included and extensible.",
    )
    add_bullet(doc, "Dense: EmbeddingFactory + FAISS/In-memory")
    add_bullet(doc, "Sparse: TF-IDF (scikit-learn)")
    add_bullet(doc, "BM25: rank-bm25")
    add_bullet(doc, "Fusion: weighted combination of normalized scores")
    add_bullet(doc, "Metadata filtering by source/page/mime_type")
    add_paragraph(doc, "Server-side now selects only the single most relevant document for citations and context.")

    # 9. Generation Pipeline
    add_heading(doc, "9. Advanced Generation Pipeline", 1)
    add_paragraph(
        doc,
        "System composes a system prompt and user prompt with retrieved context. Generation uses the selected backend. "
        "Confidence is estimated dynamically from retrieval strength (top score and margin). WebSocket route supports streaming tokens.",
    )

    # 10. Dynamic Model Adaptation
    add_heading(doc, "10. Dynamic Model Adaptation", 1)
    add_paragraph(
        doc,
        "Admin hot-swap endpoints switch embeddings, generation, and retriever backends at runtime. Embedding/retriever swaps "
        "trigger automatic reindexing from the persistent document store. The Chainlit UI provides dropdowns and a /model command to switch generation models in real time.",
    )

    # 11. Observability & Security
    add_heading(doc, "11. Observability & Security", 1)
    add_bullet(doc, "Metrics: /metrics via prometheus-fastapi-instrumentator")
    add_bullet(doc, "Logging: structured logging, level configurable")
    add_bullet(doc, "Security: response headers (X-Content-Type-Options, X-Frame-Options, Referrer-Policy)")
    add_bullet(doc, "Admin protection: X-API-Key header if admin_api_key is set")

    # 12. Testing
    add_heading(doc, "12. Testing", 1)
    add_bullet(doc, "tests/test_retrieval.py – Index & hybrid search flow")
    add_bullet(doc, "tests/test_adaptation.py – Settings hot-swap state")
    add_code_block(doc, "pytest -q")

    # 13. Performance & Tuning
    add_heading(doc, "13. Performance & Tuning", 1)
    add_bullet(doc, "Use SentenceTransformers locally for faster embeddings when API keys not available")
    add_bullet(doc, "Use FAISS for larger indices; persist vectors externally for production")
    add_bullet(doc, "Tune fusion weights and top_k per use case")
    add_bullet(doc, "Enable GPU for embeddings/generation where possible")

    # 14. Troubleshooting
    add_heading(doc, "14. Troubleshooting", 1)
    add_bullet(doc, "Form uploads require python-multipart installed")
    add_bullet(doc, "Scanned PDFs require Tesseract and poppler (for pdf2image)")
    add_bullet(doc, "Port 8000 in use: kill old Uvicorn or use a different port")
    add_bullet(doc, "Admin 401: set ADMIN_API_KEY in backend and RAG_ADMIN_API_KEY in UI env")

    # 15. Output JSON Schema
    add_heading(doc, "15. Output JSON Schema", 1)
    add_code_block(
        doc,
        """
{
  "query_id": "uuid-string",
  "answer": {
    "content": "Primary response content",
    "reasoning_steps": ["Step 1: Analysis", "Step 2: Synthesis"],
    "confidence": 0.87,
    "uncertainty_factors": ["Limited source diversity", "Temporal constraints"]
  },
  "citations": [
    {
      "document": "filename.pdf",
      "pages": [12, 13],
      "chunk_id": "chunk_001",
      "excerpt": "Relevant text snippet with context",
      "relevance_score": 0.94,
      "credibility_score": 0.89,
      "extraction_method": "dense_retrieval"
    }
  ],
  "alternative_answers": [
    {
      "content": "Alternative interpretation",
      "confidence": 0.72,
      "supporting_citations": []
    }
  ],
  "context_analysis": {
    "total_chunks_analyzed": 847,
    "retrieval_methods_used": ["dense", "sparse", "graph"],
    "cross_document_connections": 3,
    "temporal_relevance": "current"
  },
  "performance_metrics": {
    "retrieval_latency_ms": 245,
    "generation_latency_ms": 1100,
    "total_response_time_ms": 1345,
    "tokens_processed": 15420,
    "cost_estimate_usd": 0.034
  },
  "system_metadata": {
    "embedding_model": "text-embedding-3-large",
    "generation_model": "gpt-4o",
    "retrieval_strategy": "hybrid_weighted",
    "timestamp": "2025-08-11T10:30:45Z"
  }
}
        """.strip(),
    )

    return doc


def main() -> None:
    doc = build_doc()
    doc.save(OUTPUT_PATH)
    print(f"Wrote {OUTPUT_PATH}")


if __name__ == "__main__":
    main()


